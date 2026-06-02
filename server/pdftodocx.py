"""Hybrid PDF to DOCX converter.

This script is executed by the Node.js backend as a child process.
It uses layout-aware text reconstruction for simple text PDFs,
pdf2docx for more complex text layouts, and OCR for scanned PDFs.
It can also run in analyze-only mode to report PDF type metadata.
It prints a JSON result to stdout and uses exit code 2 to signal low confidence.
"""

from __future__ import annotations

import argparse
import io
import json
import logging
import os
import sys
import tempfile
import traceback
from pathlib import Path

from docx.shared import Pt

logging.basicConfig(
    level=logging.INFO,
    format="[hybrid-converter] %(levelname)s %(message)s",
    stream=sys.stderr,
)
log = logging.getLogger("hybrid")

CONFIDENCE_THRESHOLD = 0.60
SCANNED_TEXT_DENSITY_THRESHOLD = 0.01
SCANNED_TEXT_MIN_CHARS_PER_PAGE = 400
TEXT_PAGE_MIN_CHARS = 120
TEXT_PAGE_MIN_WORDS = 20
SCAN_PAGE_MAX_CHARS = 40


def _safe_float(value: object, default: float = -1.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _normalize_block_lines(text: str) -> list[str]:
    lines: list[str] = []

    for raw_line in text.replace("\xa0", " ").splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line == "|" and lines:
            lines[-1] = f"{lines[-1].rstrip()} |"
            continue

        lines.append(line)

    return lines


def analyze_pdf(pdf_path: str) -> dict:
    import fitz

    doc = fitz.open(pdf_path)
    page_count = len(doc)
    total_chars = 0
    total_words = 0
    total_image_area = 0.0
    total_page_area = 0.0
    text_pages = 0
    scanned_pages = 0
    hybrid_pages = 0
    has_tables_hint = False
    has_multi_column_hint = False

    for page in doc:
        rect = page.rect
        page_area = rect.width * rect.height
        total_page_area += page_area

        text = page.get_text("text")
        chars = len(text.strip())
        total_chars += chars
        word_count = len(page.get_text("words"))
        total_words += word_count

        density = chars / page_area if page_area > 0 else 0
        has_image_heavy_layout = False
        try:
            image_blocks = page.get_text("dict").get("blocks", [])
            for block in image_blocks:
                if block.get("type") == 1:
                    bbox = block.get("bbox", [0, 0, 0, 0])
                    image_area = max(0, bbox[2] - bbox[0]) * max(0, bbox[3] - bbox[1])
                    if page_area > 0 and image_area / page_area > 0.1:
                        has_image_heavy_layout = True
                    total_image_area += image_area
        except Exception:
            pass

        is_text_page = chars >= TEXT_PAGE_MIN_CHARS or word_count >= TEXT_PAGE_MIN_WORDS
        is_scan_page = (
            chars <= SCAN_PAGE_MAX_CHARS
            and word_count <= 8
            and (density < SCANNED_TEXT_DENSITY_THRESHOLD or has_image_heavy_layout)
        )

        if is_text_page:
            text_pages += 1
        elif is_scan_page:
            scanned_pages += 1
        else:
            hybrid_pages += 1

        try:
            drawings = page.get_drawings()
            h_lines = [item for item in drawings if abs(item["rect"].height) < 3 and item["rect"].width > 20]
            v_lines = [item for item in drawings if abs(item["rect"].width) < 3 and item["rect"].height > 20]
            if len(h_lines) >= 3 and len(v_lines) >= 2:
                has_tables_hint = True
        except Exception:
            pass

        try:
            blocks = page.get_text("blocks")
            if len(blocks) >= 4:
                x_centers = [(block[0] + block[2]) / 2 for block in blocks if block[4].strip()]
                if x_centers:
                    page_mid = rect.width / 2
                    left_blocks = sum(1 for x in x_centers if x < page_mid - 20)
                    right_blocks = sum(1 for x in x_centers if x > page_mid + 20)
                    if left_blocks >= 2 and right_blocks >= 2:
                        has_multi_column_hint = True
        except Exception:
            pass

    doc.close()

    is_scanned = scanned_pages >= max(1, int(page_count * 0.7)) and text_pages == 0
    dominant_mode = "scan" if is_scanned else ("text" if text_pages >= max(1, scanned_pages) else "hybrid")
    image_ratio = total_image_area / total_page_area if total_page_area > 0 else 0.0

    return {
        "page_count": page_count,
        "is_scanned": is_scanned,
        "dominant_mode": dominant_mode,
        "has_tables": has_tables_hint,
        "has_multi_column": has_multi_column_hint,
        "text_density": total_chars / max(page_count, 1),
        "word_density": total_words / max(page_count, 1),
        "image_ratio": round(image_ratio, 3),
        "text_pages": text_pages,
        "scanned_pages": scanned_pages,
        "hybrid_pages": hybrid_pages,
    }


def summarize_pdf_type(analysis: dict) -> dict:
    if analysis.get("is_scanned"):
        pdf_type = "scan"
        recommended_method = "ocr"
    elif analysis.get("has_tables") or analysis.get("has_multi_column") or analysis.get("hybrid_pages", 0) > 0:
        pdf_type = "text-complex"
        recommended_method = "pdf2docx"
    else:
        pdf_type = "text-simple"
        recommended_method = "layout-text"

    return {
        "pdf_type": pdf_type,
        "recommended_method": recommended_method,
    }


def convert_with_pdf2docx(pdf_path: str, output_path: str) -> dict:
    from pdf2docx import Converter

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    cv = Converter(pdf_path)
    try:
        cv.convert(output_path, start=0, end=None)
    finally:
        cv.close()

    exists = os.path.exists(output_path) and os.path.getsize(output_path) > 0
    return {
        "method": "pdf2docx",
        "success": exists,
        "confidence": 0.80 if exists else 0.0,
    }


def convert_with_layout_text(pdf_path: str, output_path: str) -> dict:
    import fitz
    from docx import Document

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc_out = Document()
    doc_out.core_properties.author = "hybrid-converter"
    style = doc_out.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    pdf_doc = fitz.open(pdf_path)
    try:
        total_pages = len(pdf_doc)
        for page_index, page in enumerate(pdf_doc):
            blocks = page.get_text("blocks", sort=True)
            page_has_content = False

            for block in blocks:
                text = str(block[4] or "").strip()
                lines = _normalize_block_lines(text)
                if not lines:
                    continue

                paragraph = doc_out.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

                for line_index, line in enumerate(lines):
                    run = paragraph.add_run(line)
                    if line_index < len(lines) - 1:
                        run.add_break()

                page_has_content = True

            if page_index < total_pages - 1 and page_has_content:
                doc_out.add_page_break()
    finally:
        pdf_doc.close()

    doc_out.save(output_path)

    exists = os.path.exists(output_path) and os.path.getsize(output_path) > 0
    return {
        "method": "layout-text",
        "success": exists,
        "confidence": 0.90 if exists else 0.0,
    }


def _configure_tesseract(lang: str) -> None:
    override = os.environ.get("TESSERACT_CMD", "").strip()
    if override:
        import pytesseract

        pytesseract.pytesseract.tesseract_cmd = override


def _extract_ocr_paragraphs(img, lang: str) -> tuple[list[str], float, int]:
    import pytesseract
    from pytesseract import Output

    custom_config = f"--psm 3 --oem 3 -l {lang}"
    data = pytesseract.image_to_data(img, config=custom_config, output_type=Output.DICT)

    grouped_lines: dict[tuple[int, int, int], list[str]] = {}
    grouped_confidences: list[float] = []

    for index, raw_text in enumerate(data.get("text", [])):
        text = str(raw_text).strip()
        if not text:
            continue

        key = (
            int(data.get("block_num", [0])[index]),
            int(data.get("par_num", [0])[index]),
            int(data.get("line_num", [0])[index]),
        )
        grouped_lines.setdefault(key, []).append(text)

        confidence_value = _safe_float(data.get("conf", ["-1"])[index])
        if confidence_value >= 0:
            grouped_confidences.append(confidence_value)

    paragraphs: list[str] = []
    paragraph_buffer: dict[tuple[int, int], list[str]] = {}

    for (block_num, par_num, line_num) in sorted(grouped_lines.keys()):
        line_text = " ".join(grouped_lines[(block_num, par_num, line_num)]).strip()
        if not line_text:
            continue
        paragraph_buffer.setdefault((block_num, par_num), []).append(line_text)

    for _, lines in paragraph_buffer.items():
        paragraph = "\n".join(lines).strip()
        if paragraph:
            paragraphs.append(paragraph)

    avg_confidence = sum(grouped_confidences) / len(grouped_confidences) if grouped_confidences else 0.0
    word_count = len(grouped_lines)
    return paragraphs, avg_confidence, word_count


def convert_with_ocr(pdf_path: str, output_path: str, lang: str = "ind+eng") -> dict:
    import fitz
    import pytesseract
    from PIL import Image
    from docx import Document
    from docx.shared import Pt

    _configure_tesseract(lang)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc_out = Document()
    doc_out.core_properties.author = "hybrid-converter"
    style = doc_out.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    pdf_doc = fitz.open(pdf_path)
    total_pages = len(pdf_doc)
    recognized_chars = 0
    ocr_confidences: list[float] = []
    detected_words = 0

    for page_num, page in enumerate(pdf_doc):
        log.info("OCR halaman %d/%d", page_num + 1, total_pages)
        matrix = fitz.Matrix(300 / 72, 300 / 72)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        img = Image.open(io.BytesIO(pix.tobytes("png")))

        paragraphs, page_confidence, page_word_count = _extract_ocr_paragraphs(img, lang)
        if not paragraphs:
            custom_config = f"--psm 3 --oem 3 -l {lang}"
            ocr_text = pytesseract.image_to_string(img, config=custom_config)
            paragraphs = [part.strip() for part in ocr_text.split("\n\n") if part.strip()]
            recognized_chars += len(ocr_text.strip())
        else:
            recognized_chars += sum(len(paragraph.strip()) for paragraph in paragraphs)

        detected_words += page_word_count
        if page_confidence > 0:
            ocr_confidences.append(page_confidence)

        if not paragraphs:
            doc_out.add_paragraph("")
        for paragraph_text in paragraphs:
            lines = paragraph_text.split("\n")
            if len(lines) == 1 and len(paragraph_text) < 60 and paragraph_text.isupper():
                doc_out.add_heading(paragraph_text, level=1)
            elif len(lines) == 1 and len(paragraph_text) < 80:
                doc_out.add_heading(paragraph_text, level=2)
            else:
                doc_out.add_paragraph(paragraph_text)

        if page_num < total_pages - 1:
            doc_out.add_page_break()

    pdf_doc.close()
    doc_out.save(output_path)

    avg_chars_per_page = recognized_chars / max(total_pages, 1)
    avg_ocr_confidence = sum(ocr_confidences) / len(ocr_confidences) if ocr_confidences else 0.0
    confidence = min(0.94, 0.30 + (avg_chars_per_page / 2200) * 0.45 + (avg_ocr_confidence / 100) * 0.25)
    return {
        "method": "ocr+python-docx",
        "success": True,
        "confidence": round(confidence, 3),
        "recognized_chars": recognized_chars,
        "recognized_words": detected_words,
        "ocr_mean_confidence": round(avg_ocr_confidence, 2),
    }


def enhance_tables_with_pdfplumber(pdf_path: str, docx_path: str) -> bool:
    try:
        import pdfplumber
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        tables_found = []
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_tables = page.extract_tables()
                if page_tables:
                    tables_found.append((page_num, page_tables))

        if not tables_found:
            return False

        doc = Document(docx_path)
        existing_table_count = len(doc.tables)
        total_pdfplumber_tables = sum(len(page_tables) for _, page_tables in tables_found)
        if existing_table_count >= total_pdfplumber_tables * 0.8:
            log.info("Tabel sudah tertangkap oleh pdf2docx (%d tabel), skip enhance", existing_table_count)
            return False

        log.info(
            "pdfplumber menemukan %d tabel, pdf2docx hanya %d -> append tabel di akhir",
            total_pdfplumber_tables,
            existing_table_count,
        )

        doc.add_page_break()
        doc.add_heading("Lampiran: Data Tabel", level=1)

        for page_num, page_tables in tables_found:
            for table_index, table_data in enumerate(page_tables):
                if not table_data or not table_data[0]:
                    continue

                paragraph = doc.add_paragraph(f"Halaman {page_num + 1}, Tabel {table_index + 1}:")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                rows = len(table_data)
                cols = max(len(row) for row in table_data)
                tbl = doc.add_table(rows=rows, cols=cols)
                tbl.style = "Table Grid"

                for row_index, row in enumerate(table_data):
                    for col_index, cell_text in enumerate(row):
                        if col_index >= cols:
                            continue
                        tbl.cell(row_index, col_index).text = str(cell_text or "")

        doc.save(docx_path)
        return True
    except Exception as error:
        log.warning("enhance_tables gagal (non-fatal): %s", error)
        return False


def compute_final_confidence(analysis: dict, conversion_result: dict, output_path: str) -> float:
    base_confidence = float(conversion_result.get("confidence", 0.0) or 0.0)

    if not os.path.exists(output_path):
        return 0.0

    output_size = os.path.getsize(output_path)
    if output_size < 1024:
        return 0.0

    if not analysis.get("is_scanned") and conversion_result.get("method") == "pdf2docx":
        base_confidence = min(1.0, base_confidence + 0.05)

    if analysis.get("is_scanned"):
        base_confidence = min(base_confidence, 0.85)

    if analysis.get("has_tables") and analysis.get("has_multi_column"):
        base_confidence = min(base_confidence, 0.75)

    return round(base_confidence, 3)


def main() -> int:
    parser = argparse.ArgumentParser(description="Hybrid PDF to DOCX converter")
    parser.add_argument("--input", required=True, help="Path ke file PDF input")
    parser.add_argument("--output", required=True, help="Path ke file DOCX output")
    parser.add_argument("--lang", default="ind+eng", help="Bahasa OCR Tesseract (default: ind+eng)")
    parser.add_argument("--force-ocr", action="store_true", help="Paksa pakai OCR meski PDF text-based")
    parser.add_argument("--analyze-only", action="store_true", help="Hanya analisis PDF tanpa melakukan konversi")
    args = parser.parse_args()

    pdf_path = args.input
    output_path = args.output

    if not os.path.exists(pdf_path):
        print(json.dumps({"error": f"File tidak ditemukan: {pdf_path}", "success": False}), file=sys.stderr)
        return 1

    try:
        log.info("Stage 1: Analisis PDF...")
        analysis = analyze_pdf(pdf_path)

        if args.analyze_only:
            pdf_summary = summarize_pdf_type(analysis)
            print(json.dumps({
                "success": True,
                "analysis": analysis,
                **pdf_summary,
            }))
            return 0

        use_ocr = args.force_ocr or (analysis["is_scanned"] and analysis["text_pages"] == 0)
        if use_ocr:
            log.info(
                "Stage 2: Jalur OCR (PDF scan terdeteksi, %d/%d halaman scanned)",
                analysis["scanned_pages"],
                analysis["page_count"],
            )
            conversion_result = convert_with_ocr(pdf_path, output_path, lang=args.lang)
        else:
            should_use_layout_text = not analysis["has_tables"] and not analysis["has_multi_column"]

            if should_use_layout_text:
                log.info("Stage 2: Jalur layout-text (PDF text-based sederhana)")
                conversion_result = convert_with_layout_text(pdf_path, output_path)
                if not conversion_result["success"]:
                    log.info("Stage 2 fallback: pdf2docx")
                    conversion_result = convert_with_pdf2docx(pdf_path, output_path)
            else:
                log.info("Stage 2: Jalur pdf2docx (layout lebih kompleks)")
                conversion_result = convert_with_pdf2docx(pdf_path, output_path)

            if analysis["has_tables"] and conversion_result["success"]:
                log.info("Stage 3: Enhance tabel dengan pdfplumber...")
                enhanced = enhance_tables_with_pdfplumber(pdf_path, output_path)
                if enhanced:
                    log.info("Stage 3: Tabel berhasil di-enhance")

        log.info("Stage 4: Quality check...")
        final_confidence = compute_final_confidence(analysis, conversion_result, output_path)

        result = {
            "success": bool(conversion_result.get("success")),
            "method": conversion_result.get("method"),
            "confidence": final_confidence,
            "analysis": analysis,
        }

        print(json.dumps(result))
        log.info("Selesai. Confidence: %.2f, Method: %s", final_confidence, conversion_result.get("method"))

        if final_confidence < CONFIDENCE_THRESHOLD:
            log.warning(
                "Confidence %.2f < threshold %.2f -> Node.js akan fallback ke LibreOffice",
                final_confidence,
                CONFIDENCE_THRESHOLD,
            )
            return 2

        return 0
    except ImportError as error:
        error_msg = f"Library tidak tersedia: {error}. Pastikan requirements sudah diinstall."
        log.error(error_msg)
        print(json.dumps({"error": error_msg, "success": False}))
        return 1
    except Exception as error:
        log.error("Error tidak terduga: %s\n%s", error, traceback.format_exc())
        print(json.dumps({"error": str(error), "success": False}))
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
