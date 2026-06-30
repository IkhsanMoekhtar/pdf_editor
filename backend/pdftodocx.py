# -*- coding: utf-8 -*-
"""Faithful PDF to DOCX converter - versi custom.

Strategi konversi bertingkat:
  1. FaithfulText  — untuk PDF text-based sederhana:
       Membaca span-level teks (font, ukuran, warna, bold/italic),
       mendeteksi heading, bullet, indentasi, gambar inline, dan tabel
       menggunakan pdfplumber+PyMuPDF, lalu menulis ke DOCX dengan
       python-docx sehingga output semirip mungkin dengan tampilan asli PDF.
  2. pdf2docx     — fallback untuk layout multi-kolom / kompleks.
  3. OCR          — fallback untuk PDF hasil scan.

Script ini dipanggil sebagai child-process oleh Node.js backend.
Hasil JSON dicetak ke stdout; exit code 2 berarti confidence rendah.
"""

from __future__ import annotations

import argparse
import io
import json
import logging
import os
import re
import sys
import tempfile
import traceback
from collections import defaultdict
from pathlib import Path
from typing import Any

logging.basicConfig(
    level=logging.INFO,
    format="[pdf-faithful] %(levelname)s %(message)s",
    stream=sys.stderr,
)
log = logging.getLogger("faithful")

# ─── Konstanta ────────────────────────────────────────────────────────────────
CONFIDENCE_THRESHOLD = 0.60
SCANNED_TEXT_DENSITY_THRESHOLD = 0.01
SCANNED_TEXT_MIN_CHARS_PER_PAGE = 400
TEXT_PAGE_MIN_CHARS = 120
TEXT_PAGE_MIN_WORDS = 20
SCAN_PAGE_MAX_CHARS = 40

# Toleransi pengelompokan baris (dalam poin PDF)
LINE_Y_TOLERANCE = 3.0
PARA_BREAK_THRESHOLD = 1.4   # lipatan line-height untuk deteksi paragraf baru
INDENT_THRESHOLD = 12.0       # minimum selisih x untuk dianggap indent
HEADING_SCALE_MIN = 1.15      # font size ≥ 115 % rata-rata → heading
BULLET_CHARS = set("•●○◦▪▸►▶‣⁃")


# ─── Helper ───────────────────────────────────────────────────────────────────

def _safe_float(value: object, default: float = -1.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _rgb_from_int(color_int: int | None) -> tuple[int, int, int] | None:
    """Ubah integer warna PyMuPDF (0xRRGGBB) ke tuple RGB."""
    if color_int is None:
        return None
    r = (color_int >> 16) & 0xFF
    g = (color_int >> 8) & 0xFF
    b = color_int & 0xFF
    return (r, g, b)


def _rgb_from_tuple(color_tuple) -> tuple[int, int, int] | None:
    """Ubah tuple float (0.0–1.0) ke tuple int RGB."""
    if not color_tuple:
        return None
    try:
        if isinstance(color_tuple, (int, float)):
            v = int(color_tuple * 255)
            return (v, v, v)
        if len(color_tuple) == 3:
            return tuple(int(c * 255) for c in color_tuple)
        if len(color_tuple) == 4:
            # CMYK → RGB sederhana
            c, m, y, k = [float(x) for x in color_tuple]
            r = int(255 * (1 - c) * (1 - k))
            g = int(255 * (1 - m) * (1 - k))
            b = int(255 * (1 - y) * (1 - k))
            return (r, g, b)
    except Exception:
        pass
    return None


# ─── Analisis PDF ─────────────────────────────────────────────────────────────

def analyze_pdf(pdf_path: str) -> dict:
    import fitz

    doc = fitz.open(pdf_path)
    page_count = len(doc)
    total_chars = 0
    total_words = 0
    total_image_area = 0.0
    total_page_area = 0.0
    text_pages = 0
    scanned_pages = 0
    hybrid_pages = 0
    has_tables_hint = False
    has_multi_column_hint = False

    for page in doc:
        rect = page.rect
        page_area = rect.width * rect.height
        total_page_area += page_area

        text = page.get_text("text")
        chars = len(text.strip())
        total_chars += chars
        word_count = len(page.get_text("words"))
        total_words += word_count

        density = chars / page_area if page_area > 0 else 0
        has_image_heavy_layout = False
        try:
            image_blocks = page.get_text("dict").get("blocks", [])
            for block in image_blocks:
                if block.get("type") == 1:
                    bbox = block.get("bbox", [0, 0, 0, 0])
                    image_area = max(0, bbox[2] - bbox[0]) * max(0, bbox[3] - bbox[1])
                    if page_area > 0 and image_area / page_area > 0.1:
                        has_image_heavy_layout = True
                    total_image_area += image_area
        except Exception:
            pass

        is_text_page = chars >= TEXT_PAGE_MIN_CHARS or word_count >= TEXT_PAGE_MIN_WORDS
        is_scan_page = (
            chars <= SCAN_PAGE_MAX_CHARS
            and word_count <= 8
            and (density < SCANNED_TEXT_DENSITY_THRESHOLD or has_image_heavy_layout)
        )

        if is_text_page:
            text_pages += 1
        elif is_scan_page:
            scanned_pages += 1
        else:
            hybrid_pages += 1

        try:
            drawings = page.get_drawings()
            h_lines = [d for d in drawings if abs(d["rect"].height) < 3 and d["rect"].width > 20]
            v_lines = [d for d in drawings if abs(d["rect"].width) < 3 and d["rect"].height > 20]
            if len(h_lines) >= 3 and len(v_lines) >= 2:
                has_tables_hint = True
        except Exception:
            pass

        try:
            blocks = page.get_text("blocks")
            if len(blocks) >= 4:
                x_centers = [(b[0] + b[2]) / 2 for b in blocks if str(b[4]).strip()]
                if x_centers:
                    page_mid = rect.width / 2
                    left_blocks = sum(1 for x in x_centers if x < page_mid - 20)
                    right_blocks = sum(1 for x in x_centers if x > page_mid + 20)
                    if left_blocks >= 2 and right_blocks >= 2:
                        has_multi_column_hint = True
        except Exception:
            pass

    doc.close()

    is_scanned = scanned_pages >= max(1, int(page_count * 0.7)) and text_pages == 0
    dominant_mode = "scan" if is_scanned else ("text" if text_pages >= max(1, scanned_pages) else "hybrid")
    image_ratio = total_image_area / total_page_area if total_page_area > 0 else 0.0

    return {
        "page_count": page_count,
        "is_scanned": is_scanned,
        "dominant_mode": dominant_mode,
        "has_tables": has_tables_hint,
        "has_multi_column": has_multi_column_hint,
        "text_density": total_chars / max(page_count, 1),
        "word_density": total_words / max(page_count, 1),
        "image_ratio": round(image_ratio, 3),
        "text_pages": text_pages,
        "scanned_pages": scanned_pages,
        "hybrid_pages": hybrid_pages,
    }


def summarize_pdf_type(analysis: dict) -> dict:
    if analysis.get("is_scanned"):
        pdf_type = "scan"
        recommended_method = "ocr"
    elif analysis.get("has_tables") or analysis.get("has_multi_column") or analysis.get("hybrid_pages", 0) > 0:
        pdf_type = "text-complex"
        recommended_method = "pdf2docx"
    else:
        pdf_type = "text-simple"
        recommended_method = "faithful-text"

    return {
        "pdf_type": pdf_type,
        "recommended_method": recommended_method,
    }


# ─── Konversi Faithful (custom) ───────────────────────────────────────────────

class _Span:
    """Satu potongan teks dalam baris dengan properti gaya penuh."""
    __slots__ = ("text", "font", "size", "bold", "italic", "color", "x0", "x1", "y0", "y1")

    def __init__(self, raw: dict):
        flags = int(raw.get("flags", 0))
        self.text: str = raw.get("text", "")
        self.font: str = raw.get("font", "Arial")
        self.size: float = float(raw.get("size", 11))
        self.bold: bool = bool(flags & 2**4)
        self.italic: bool = bool(flags & 2**1)
        self.color: tuple[int, int, int] | None = _rgb_from_int(raw.get("color"))
        bbox = raw.get("bbox", (0, 0, 0, 0))
        self.x0 = float(bbox[0])
        self.y0 = float(bbox[1])
        self.x1 = float(bbox[2])
        self.y1 = float(bbox[3])


class _Line:
    """Satu baris teks (kumpulan span)."""
    __slots__ = ("spans", "x0", "y0", "x1", "y1")

    def __init__(self, raw: dict):
        self.spans: list[_Span] = [_Span(s) for s in raw.get("spans", []) if str(s.get("text", "")).strip()]
        bbox = raw.get("bbox", (0, 0, 0, 0))
        self.x0 = float(bbox[0])
        self.y0 = float(bbox[1])
        self.x1 = float(bbox[2])
        self.y1 = float(bbox[3])

    @property
    def height(self) -> float:
        return self.y1 - self.y0

    @property
    def text(self) -> str:
        return "".join(s.text for s in self.spans)

    @property
    def dominant_size(self) -> float:
        sizes = [s.size for s in self.spans if s.text.strip()]
        return max(sizes) if sizes else 11.0


def _collect_page_lines(page_dict: dict) -> list[_Line]:
    """Kumpulkan semua baris dari dict halaman (hanya blok teks)."""
    lines: list[_Line] = []
    for block in page_dict.get("blocks", []):
        if block.get("type") != 0:
            continue
        for raw_line in block.get("lines", []):
            line = _Line(raw_line)
            if line.spans:
                lines.append(line)
    lines.sort(key=lambda l: (l.y0, l.x0))
    return lines


def _group_lines_into_paragraphs(lines: list[_Line]) -> list[list[_Line]]:
    """Kelompokkan baris menjadi paragraf berdasarkan jarak vertikal."""
    if not lines:
        return []

    paragraphs: list[list[_Line]] = []
    current: list[_Line] = [lines[0]]

    for prev, curr in zip(lines, lines[1:]):
        gap = curr.y0 - prev.y1
        avg_line_h = (prev.height + curr.height) / 2.0 if (prev.height + curr.height) > 0 else 12.0
        # Paragraf baru jika: gap > faktor * tinggi rata-rata, ATAU indentasi berubah jauh
        if gap > PARA_BREAK_THRESHOLD * avg_line_h or gap > 18:
            paragraphs.append(current)
            current = [curr]
        else:
            current.append(curr)

    paragraphs.append(current)
    return paragraphs


def _median_font_size(all_lines: list[_Line]) -> float:
    sizes = []
    for line in all_lines:
        for span in line.spans:
            if span.text.strip():
                sizes.append(span.size)
    if not sizes:
        return 11.0
    sizes.sort()
    mid = len(sizes) // 2
    return sizes[mid]


def _detect_bullet(text: str) -> str | None:
    """Kembalikan karakter bullet jika baris adalah list item."""
    stripped = text.lstrip()
    if stripped and stripped[0] in BULLET_CHARS:
        return stripped[0]
    # Numbered list: "1." / "1)" / "(1)"
    if re.match(r"^\d+[\.\)]\s", stripped) or re.match(r"^\(\d+\)\s", stripped):
        return "•"
    return None


def _font_family(font_name: str) -> str:
    """Ekstrak nama keluarga font yang aman untuk DOCX."""
    # Hapus subset prefix seperti "ABCDEF+"
    name = re.sub(r"^[A-Z]{6}\+", "", font_name)
    # Petakan nama umum
    lname = name.lower()
    if "arial" in lname:
        return "Arial"
    if "times" in lname or "roman" in lname:
        return "Times New Roman"
    if "calibri" in lname:
        return "Calibri"
    if "helvetica" in lname:
        return "Helvetica"
    if "courier" in lname:
        return "Courier New"
    if "verdana" in lname:
        return "Verdana"
    if "georgia" in lname:
        return "Georgia"
    if "trebuchet" in lname:
        return "Trebuchet MS"
    if "garamond" in lname:
        return "Garamond"
    if "bookman" in lname:
        return "Book Antiqua"
    # Default: gunakan nama asli jika masuk akal, else Arial
    clean = re.sub(r"[^A-Za-z0-9 \-]", "", name).strip()
    return clean if clean else "Arial"


def _add_run_with_style(paragraph, span: _Span, text: str | None = None) -> None:
    """Tambahkan run ke paragraf dengan semua properti gaya dari span."""
    from docx.shared import Pt, RGBColor

    run_text = text if text is not None else span.text
    if not run_text:
        return

    run = paragraph.add_run(run_text)
    run.font.name = _font_family(span.font)
    run.font.size = Pt(round(span.size, 1))
    run.font.bold = span.bold
    run.font.italic = span.italic

    if span.color and span.color != (0, 0, 0):
        run.font.color.rgb = RGBColor(*span.color)


def _write_paragraph_lines(
    doc_out,
    para_lines: list[_Line],
    median_size: float,
    page_width_pt: float,
    left_margin_pt: float,
) -> None:
    """Tulis satu paragraf (kumpulan baris) ke dokumen DOCX."""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not para_lines:
        return

    # Ambil baris pertama untuk menentukan properti paragraf
    first_line = para_lines[0]
    all_spans = [s for ln in para_lines for s in ln.spans if s.text.strip()]
    if not all_spans:
        return

    # Deteksi heading berdasarkan ukuran font dominan
    dominant_size = max(s.size for s in all_spans) if all_spans else 11.0
    is_heading = dominant_size >= median_size * HEADING_SCALE_MIN and len("".join(s.text for s in all_spans)) < 200

    # Deteksi bullet
    first_text = "".join(s.text for s in all_spans)
    bullet_char = _detect_bullet(first_text)

    # Deteksi alignment berdasarkan posisi x di halaman
    usable_width = page_width_pt - left_margin_pt * 2
    text_start = first_line.x0
    text_end = max(ln.x1 for ln in para_lines)
    text_center = (text_start + text_end) / 2
    page_center = page_width_pt / 2

    if abs(text_center - page_center) < 30 and (text_end - text_start) < usable_width * 0.6:
        alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif text_end > page_width_pt - left_margin_pt - 20 and text_start > left_margin_pt + INDENT_THRESHOLD:
        alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Hitung indentasi kiri (dalam inci)
    indent_pt = max(0.0, text_start - left_margin_pt)
    indent_inches = indent_pt / 72.0

    # Buat paragraf
    if is_heading:
        level = 1 if dominant_size >= median_size * 1.5 else 2
        para = doc_out.add_heading("", level=level)
    else:
        para = doc_out.add_paragraph()

    para.alignment = alignment
    if indent_inches > 0.1:
        para.paragraph_format.left_indent = Inches(min(indent_inches, 3.0))

    # Spasi vertikal
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)

    # Tambahkan konten per baris
    for line_idx, line in enumerate(para_lines):
        if line_idx > 0:
            # Tambahkan line break antar baris dalam satu paragraf
            if para.runs:
                para.runs[-1].add_break()

        for span in line.spans:
            if not span.text:
                continue
            # Hapus bullet dari teks pertama jika terdeteksi
            if line_idx == 0 and bullet_char and span.text.lstrip().startswith(tuple(BULLET_CHARS)):
                cleaned = span.text.lstrip()
                # Hapus karakter bullet di awal
                cleaned = re.sub(r"^[•●○◦▪▸►▶‣⁃]\s*", "", cleaned)
                _add_run_with_style(para, span, cleaned)
            else:
                _add_run_with_style(para, span)


def _embed_page_images(doc_out, page, page_index: int, page_count: int, temp_dir: str) -> None:
    """Ekstrak dan embed gambar dari halaman PDF ke DOCX."""
    from docx.shared import Inches

    try:
        image_list = page.get_images(full=True)
        for img_index, img_info in enumerate(image_list):
            xref = img_info[0]
            try:
                base_image = page.parent.extract_image(xref)
                img_bytes = base_image.get("image")
                img_ext = base_image.get("ext", "png")
                if not img_bytes:
                    continue

                img_path = os.path.join(temp_dir, f"page{page_index}_img{img_index}.{img_ext}")
                with open(img_path, "wb") as fh:
                    fh.write(img_bytes)

                # Tentukan lebar gambar di dokumen
                width_pt = float(base_image.get("width", 400))
                height_pt = float(base_image.get("height", 300))
                # Skala agar muat dalam halaman (maks 6 inci = 432 pt)
                max_width_inches = 6.0
                width_inches = min(width_pt / 96.0, max_width_inches)

                para = doc_out.add_paragraph()
                para.alignment = 1  # CENTER
                run = para.add_run()
                run.add_picture(img_path, width=Inches(width_inches))
            except Exception as img_err:
                log.debug("Gambar %d halaman %d dilewati: %s", img_index, page_index, img_err)
    except Exception as err:
        log.debug("extract_images halaman %d: %s", page_index, err)


def _embed_tables_pdfplumber(doc_out, pdf_path: str, page_num: int) -> int:
    """Embed tabel dari halaman tertentu menggunakan pdfplumber. Kembalikan jumlah tabel."""
    try:
        import pdfplumber

        with pdfplumber.open(pdf_path) as pdf:
            if page_num >= len(pdf.pages):
                return 0
            plumber_page = pdf.pages[page_num]
            tables = plumber_page.extract_tables()

        count = 0
        for table_data in tables:
            if not table_data or not table_data[0]:
                continue
            rows = len(table_data)
            cols = max(len(row) for row in table_data)
            if rows < 1 or cols < 1:
                continue

            tbl = doc_out.add_table(rows=rows, cols=cols)
            tbl.style = "Table Grid"
            for row_i, row in enumerate(table_data):
                for col_i, cell_val in enumerate(row):
                    if col_i < cols:
                        tbl.cell(row_i, col_i).text = str(cell_val or "")
            doc_out.add_paragraph()  # Spasi setelah tabel
            count += 1

        return count
    except Exception as err:
        log.debug("pdfplumber halaman %d: %s", page_num, err)
        return 0


def convert_with_faithful_text(pdf_path: str, output_path: str) -> dict:
    """
    Konversi PDF ke DOCX dengan preservasi gaya teks maksimal.
    Membaca font, ukuran, bold, italic, warna, posisi, gambar, dan tabel.
    """
    import fitz
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc_out = Document()
    # Hapus paragraf default kosong
    for para in doc_out.paragraphs:
        p = para._element
        p.getparent().remove(p)

    pdf_doc = fitz.open(pdf_path)
    total_pages = len(pdf_doc)

    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            for page_index, page in enumerate(pdf_doc):
                # ── 1. Ukuran halaman ───────────────────────────────────────
                rect = page.rect
                page_w_pt = rect.width   # dalam poin PDF (1 pt = 1/72 inci)
                page_h_pt = rect.height

                # Setel ukuran halaman dokumen dari halaman pertama
                if page_index == 0:
                    section = doc_out.sections[0]
                    section.page_width = int(page_w_pt * 12700)   # EMU (1 pt = 12700 EMU)
                    section.page_height = int(page_h_pt * 12700)
                    # Margin kira-kira 1 inci = 914400 EMU
                    margin_emu = int(72 * 12700)   # 72 pt ≈ 1 inci
                    section.left_margin = margin_emu
                    section.right_margin = margin_emu
                    section.top_margin = margin_emu
                    section.bottom_margin = margin_emu
                    left_margin_pt = 72.0
                else:
                    section = doc_out.sections[0]
                    left_margin_pt = section.left_margin / 12700

                # ── 2. Tabel dengan pdfplumber ──────────────────────────────
                table_count = _embed_tables_pdfplumber(doc_out, pdf_path, page_index)
                if table_count > 0:
                    log.info("Halaman %d: %d tabel diembed", page_index + 1, table_count)

                # ── 3. Gambar inline ────────────────────────────────────────
                _embed_page_images(doc_out, page, page_index, total_pages, temp_dir)

                # ── 4. Teks dengan gaya ─────────────────────────────────────
                page_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE | fitz.TEXT_MEDIABOX_CLIP)
                all_lines = _collect_page_lines(page_dict)
                median_size = _median_font_size(all_lines)

                para_groups = _group_lines_into_paragraphs(all_lines)
                for para_lines in para_groups:
                    _write_paragraph_lines(
                        doc_out,
                        para_lines,
                        median_size,
                        page_w_pt,
                        left_margin_pt,
                    )

                # ── 5. Page break ───────────────────────────────────────────
                if page_index < total_pages - 1:
                    doc_out.add_page_break()

        finally:
            pdf_doc.close()

    doc_out.save(output_path)

    exists = os.path.exists(output_path) and os.path.getsize(output_path) > 0
    return {
        "method": "faithful-text",
        "success": exists,
        "confidence": 0.88 if exists else 0.0,
    }


# ─── Konversi pdf2docx ────────────────────────────────────────────────────────

def convert_with_pdf2docx(pdf_path: str, output_path: str) -> dict:
    from pdf2docx import Converter

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    cv = Converter(pdf_path)
    try:
        cv.convert(
            output_path,
            start=0,
            end=None,
            word_margin=0.3,
            char_margin=3.0,
            line_margin=0.3,
        )
    finally:
        cv.close()

    exists = os.path.exists(output_path) and os.path.getsize(output_path) > 0
    return {
        "method": "pdf2docx",
        "success": exists,
        "confidence": 0.80 if exists else 0.0,
    }


# ─── Konversi OCR ─────────────────────────────────────────────────────────────

def _configure_tesseract(lang: str) -> None:
    override = os.environ.get("TESSERACT_CMD", "").strip()
    if override:
        import pytesseract
        pytesseract.pytesseract.tesseract_cmd = override


def _extract_ocr_paragraphs(img, lang: str) -> tuple[list[str], float, int]:
    import pytesseract
    from pytesseract import Output

    custom_config = f"--psm 3 --oem 3 -l {lang}"
    data = pytesseract.image_to_data(img, config=custom_config, output_type=Output.DICT)

    grouped_lines: dict[tuple[int, int, int], list[str]] = {}
    grouped_confidences: list[float] = []

    for index, raw_text in enumerate(data.get("text", [])):
        text = str(raw_text).strip()
        if not text:
            continue
        key = (
            int(data.get("block_num", [0])[index]),
            int(data.get("par_num", [0])[index]),
            int(data.get("line_num", [0])[index]),
        )
        grouped_lines.setdefault(key, []).append(text)
        confidence_value = _safe_float(data.get("conf", ["-1"])[index])
        if confidence_value >= 0:
            grouped_confidences.append(confidence_value)

    paragraph_buffer: dict[tuple[int, int], list[str]] = {}
    for (block_num, par_num, line_num) in sorted(grouped_lines.keys()):
        line_text = " ".join(grouped_lines[(block_num, par_num, line_num)]).strip()
        if line_text:
            paragraph_buffer.setdefault((block_num, par_num), []).append(line_text)

    paragraphs: list[str] = []
    for _, lines in paragraph_buffer.items():
        paragraph = "\n".join(lines).strip()
        if paragraph:
            paragraphs.append(paragraph)

    avg_confidence = sum(grouped_confidences) / len(grouped_confidences) if grouped_confidences else 0.0
    word_count = len(grouped_lines)
    return paragraphs, avg_confidence, word_count


def convert_with_ocr(pdf_path: str, output_path: str, lang: str = "ind+eng") -> dict:
    import fitz
    import pytesseract
    from PIL import Image
    from docx import Document
    from docx.shared import Pt

    _configure_tesseract(lang)
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc_out = Document()
    doc_out.core_properties.author = "pdf-faithful-converter"
    style = doc_out.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    pdf_doc = fitz.open(pdf_path)
    total_pages = len(pdf_doc)
    recognized_chars = 0
    ocr_confidences: list[float] = []
    detected_words = 0

    for page_num, page in enumerate(pdf_doc):
        log.info("OCR halaman %d/%d", page_num + 1, total_pages)
        matrix = fitz.Matrix(300 / 72, 300 / 72)
        pix = page.get_pixmap(matrix=matrix, alpha=False)
        img = Image.open(io.BytesIO(pix.tobytes("png")))

        paragraphs, page_confidence, page_word_count = _extract_ocr_paragraphs(img, lang)
        if not paragraphs:
            custom_config = f"--psm 3 --oem 3 -l {lang}"
            ocr_text = pytesseract.image_to_string(img, config=custom_config)
            paragraphs = [part.strip() for part in ocr_text.split("\n\n") if part.strip()]
            recognized_chars += len(ocr_text.strip())
        else:
            recognized_chars += sum(len(p.strip()) for p in paragraphs)

        detected_words += page_word_count
        if page_confidence > 0:
            ocr_confidences.append(page_confidence)

        if not paragraphs:
            doc_out.add_paragraph("")
        for para_text in paragraphs:
            lines = para_text.split("\n")
            if len(lines) == 1 and len(para_text) < 60 and para_text.isupper():
                doc_out.add_heading(para_text, level=1)
            elif len(lines) == 1 and len(para_text) < 80:
                doc_out.add_heading(para_text, level=2)
            else:
                doc_out.add_paragraph(para_text)

        if page_num < total_pages - 1:
            doc_out.add_page_break()

    pdf_doc.close()
    doc_out.save(output_path)

    avg_chars = recognized_chars / max(total_pages, 1)
    avg_conf = sum(ocr_confidences) / len(ocr_confidences) if ocr_confidences else 0.0
    confidence = min(0.94, 0.30 + (avg_chars / 2200) * 0.45 + (avg_conf / 100) * 0.25)
    return {
        "method": "ocr+python-docx",
        "success": True,
        "confidence": round(confidence, 3),
        "recognized_chars": recognized_chars,
        "recognized_words": detected_words,
        "ocr_mean_confidence": round(avg_conf, 2),
    }


# ─── Enhance tabel dengan pdfplumber ─────────────────────────────────────────

def enhance_tables_with_pdfplumber(pdf_path: str, docx_path: str) -> bool:
    """
    Jika pdf2docx tidak menangkap semua tabel, append tabel tambahan di akhir dokumen.
    """
    try:
        import pdfplumber
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        tables_found = []
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_tables = page.extract_tables()
                if page_tables:
                    tables_found.append((page_num, page_tables))

        if not tables_found:
            return False

        doc = Document(docx_path)
        existing_table_count = len(doc.tables)
        total_pdfplumber_tables = sum(len(pt) for _, pt in tables_found)
        if existing_table_count >= total_pdfplumber_tables * 0.8:
            log.info("Tabel sudah tertangkap (%d tabel), skip enhance", existing_table_count)
            return False

        log.info(
            "pdfplumber: %d tabel, pdf2docx: %d → append",
            total_pdfplumber_tables,
            existing_table_count,
        )

        doc.add_page_break()
        doc.add_heading("Lampiran: Data Tabel", level=1)

        for page_num, page_tables in tables_found:
            for table_index, table_data in enumerate(page_tables):
                if not table_data or not table_data[0]:
                    continue

                para = doc.add_paragraph(f"Halaman {page_num + 1}, Tabel {table_index + 1}:")
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

                rows = len(table_data)
                cols = max(len(row) for row in table_data)
                tbl = doc.add_table(rows=rows, cols=cols)
                tbl.style = "Table Grid"

                for row_i, row in enumerate(table_data):
                    for col_i, cell_val in enumerate(row):
                        if col_i < cols:
                            tbl.cell(row_i, col_i).text = str(cell_val or "")

        doc.save(docx_path)
        return True
    except Exception as error:
        log.warning("enhance_tables gagal (non-fatal): %s", error)
        return False


# ─── Quality check ────────────────────────────────────────────────────────────

def compute_final_confidence(analysis: dict, conversion_result: dict, output_path: str) -> float:
    base_confidence = float(conversion_result.get("confidence", 0.0) or 0.0)

    if not os.path.exists(output_path):
        return 0.0

    output_size = os.path.getsize(output_path)
    if output_size < 1024:
        return 0.0

    method = conversion_result.get("method", "")
    if not analysis.get("is_scanned") and method in ("pdf2docx", "faithful-text"):
        base_confidence = min(1.0, base_confidence + 0.05)

    if analysis.get("is_scanned"):
        base_confidence = min(base_confidence, 0.85)

    if analysis.get("has_tables") and analysis.get("has_multi_column"):
        base_confidence = min(base_confidence, 0.75)

    return round(base_confidence, 3)


# ─── Main ─────────────────────────────────────────────────────────────────────

def main() -> int:
    parser = argparse.ArgumentParser(description="Faithful PDF to DOCX converter")
    parser.add_argument("--input", required=True, help="Path ke file PDF input")
    parser.add_argument("--output", required=True, help="Path ke file DOCX output")
    parser.add_argument("--lang", default="ind+eng", help="Bahasa OCR Tesseract (default: ind+eng)")
    parser.add_argument("--force-ocr", action="store_true", help="Paksa pakai OCR meski PDF text-based")
    parser.add_argument("--analyze-only", action="store_true", help="Hanya analisis PDF tanpa konversi")
    args = parser.parse_args()

    pdf_path = args.input
    output_path = args.output

    # Validasi path output harus berada dalam direktori temporary sistem.
    import tempfile as _tmpmod
    _abs_output = os.path.realpath(os.path.abspath(output_path))
    _tmp_root = os.path.realpath(_tmpmod.gettempdir())
    if not (_abs_output.startswith(_tmp_root + os.sep) or _abs_output == _tmp_root):
        print(json.dumps({"error": "Path output tidak valid: harus berada dalam direktori temporary.", "success": False}))
        return 1

    if not os.path.exists(pdf_path):
        print(json.dumps({"error": f"File tidak ditemukan: {pdf_path}", "success": False}), file=sys.stderr)
        return 1


    try:
        log.info("Stage 1: Analisis PDF ...")
        analysis = analyze_pdf(pdf_path)

        if args.analyze_only:
            pdf_summary = summarize_pdf_type(analysis)
            print(json.dumps({
                "success": True,
                "analysis": analysis,
                **pdf_summary,
            }))
            return 0

        use_ocr = args.force_ocr or (analysis["is_scanned"] and analysis["text_pages"] == 0)

        if use_ocr:
            log.info(
                "Stage 2: Jalur OCR (PDF scan, %d/%d halaman scanned)",
                analysis["scanned_pages"],
                analysis["page_count"],
            )
            conversion_result = convert_with_ocr(pdf_path, output_path, lang=args.lang)

        elif analysis.get("has_multi_column") or analysis.get("hybrid_pages", 0) > 2:
            # Multi-kolom / layout sangat kompleks → pdf2docx lebih andal
            log.info("Stage 2: Jalur pdf2docx (layout kompleks / multi-kolom)")
            conversion_result = convert_with_pdf2docx(pdf_path, output_path)

            if analysis.get("has_tables") and conversion_result["success"]:
                log.info("Stage 3: Enhance tabel dengan pdfplumber ...")
                if enhance_tables_with_pdfplumber(pdf_path, output_path):
                    log.info("Stage 3: Tabel berhasil di-enhance")

        else:
            # Text-based (sederhana / berstruktur): gunakan faithful-text dulu
            log.info("Stage 2: Jalur faithful-text (preservasi gaya maksimal)")
            conversion_result = convert_with_faithful_text(pdf_path, output_path)

            if not conversion_result["success"]:
                log.info("Stage 2 fallback: pdf2docx")
                conversion_result = convert_with_pdf2docx(pdf_path, output_path)

            if analysis.get("has_tables") and conversion_result["success"]:
                log.info("Stage 3: Enhance tabel dengan pdfplumber ...")
                if enhance_tables_with_pdfplumber(pdf_path, output_path):
                    log.info("Stage 3: Tabel berhasil di-enhance")

        log.info("Stage 4: Quality check ...")
        final_confidence = compute_final_confidence(analysis, conversion_result, output_path)

        result = {
            "success": bool(conversion_result.get("success")),
            "method": conversion_result.get("method"),
            "confidence": final_confidence,
            "analysis": analysis,
        }

        print(json.dumps(result))
        log.info(
            "Selesai. Confidence: %.2f, Method: %s",
            final_confidence,
            conversion_result.get("method"),
        )

        if final_confidence < CONFIDENCE_THRESHOLD:
            log.warning(
                "Confidence %.2f < threshold %.2f → Node.js akan fallback ke LibreOffice",
                final_confidence,
                CONFIDENCE_THRESHOLD,
            )
            return 2

        return 0

    except ImportError as error:
        error_msg = f"Library tidak tersedia: {error}. Pastikan requirements sudah diinstall."
        log.error(error_msg)
        print(json.dumps({"error": error_msg, "success": False}))
        return 1
    except Exception as error:
        log.error("Error tidak terduga: %s\n%s", error, traceback.format_exc())
        print(json.dumps({"error": str(error), "success": False}))
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
